"""
📜 Sanskrit Document Retrieval (RAG) - Final Streamlit Application
This application converts the Jupyter notebook to an interactive Streamlit app
with support for multiple file types (DOCX, PDF, TXT)
"""

import re
import warnings
import os
import tempfile
from typing import List, Tuple
import streamlit as st
import torch
from dotenv import load_dotenv
import langchain 
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.prompts import PromptTemplate, ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from transformers import AutoTokenizer, AutoModelForCausalLM

# Suppress warnings
warnings.filterwarnings("ignore")
load_dotenv()

# ============================================================================
# Page Configuration
# ============================================================================
st.set_page_config(
    page_title="Sanskrit RAG System",
    page_icon="📜",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("📜 Sanskrit Document Retrieval (RAG) System")
st.markdown("""
This application extracts and indexes documents to enable Sanskrit-based 
semantic search and answer generation using Retrieval-Augmented Generation (RAG).
""")

# ============================================================================
# Helper Functions
# ============================================================================

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def is_valid_sanskrit_query(text: str, min_chars: int = 3) -> bool:
    """Validate if the query is in Sanskrit (Devanagari script)"""
    chars = re.findall(r"[\u0900-\u097F]", text)
    non_sanskrit = re.sub(r"[\u0900-\u097F\s।॥?]", "", text)
    return len(chars) >= min_chars and non_sanskrit.strip() == ""


def detect_query_type(text: str) -> str:
    """Detect whether the query is Sanskrit (Devanagari) or unknown"""
    if re.search(r"[\u0900-\u097F]", text):
        return "sanskrit_devanagari"
    if re.search(r"[a-zA-Z]", text):
        return "English"
    return "unknown"


def is_valid_query(text: str, min_chars: int = 3) -> bool:
    """Allow English and Sanskrit (Devanagari)"""
    text = text.strip()
    if len(text) < min_chars:
        return False
    return True


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        st.error("PyPDF2 library not installed. Please install: pip install PyPDF2")
        return ""
    except Exception as e:
        st.error(f"Error reading PDF file: {str(e)}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        st.error(f"Error reading TXT file: {str(e)}")
        return ""


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text based on file type"""
    if file_type.lower() == "docx":
        return extract_text_from_docx(file_path)
    elif file_type.lower() == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type.lower() == "txt":
        return extract_text_from_txt(file_path)
    else:
        st.error(f"Unsupported file type: {file_type}")
        return ""


def process_documents(raw_documents: List[LangchainDocument]) -> Tuple[List[LangchainDocument], int]:
    """Clean and process documents"""
    cleaned_docs = []
    for doc in raw_documents:
        cleaned = clean_text(doc.page_content)
        if len(cleaned) > 50:
            doc.page_content = cleaned
            cleaned_docs.append(doc)
    return cleaned_docs, len(cleaned_docs)


def create_chunks(cleaned_docs: List[LangchainDocument]) -> List[LangchainDocument]:
    """Split documents into chunks"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=150,
        separators=["॥", "।", "\n\n", "\n", " "]
    )
    documents = splitter.split_documents(cleaned_docs)
    return documents


def create_vector_store(documents: List[LangchainDocument]):
    """Create FAISS vector store"""
    import os
    os.environ["TOKENIZERS_PARALLELISM"] = "false"
    
    try:
        # Use a simple embedding approach without device specs to avoid meta tensor issues
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            cache_folder=r"models"
        )
        vectorstore = FAISS.from_documents(documents, embeddings)
        return vectorstore, embeddings
    except Exception as e:
        st.error(f"Embedding error: {str(e)}")
        st.info("Trying alternative embedding method...")
        # Fallback: create a mock vectorstore
        raise


def generate_answer(query: str, docs: List[LangchainDocument], tokenizer, model) -> str:
    """Generate answer based on retrieved documents using chat prompt"""
    context = "\n\n".join(d.page_content for d in docs)

    chat_prompt = ChatPromptTemplate.from_messages([
        ("system",
     "त्वं संस्कृत-प्रश्नोत्तर-सहायकः असि।\n"
     "नियमाः:\n"
     "१) अधोलिखित-सन्दर्भे एव आधारित्य उत्तरं लिख।\n"
     "२) सन्दर्भात् बहिः किमपि न लिख।\n"
     "३) व्याख्यां न लिख। पुनरुक्तिं न लिख।\n"
     "४) उत्तरं केवलं एकेन वाक्येन लिख।\n"
     "५) उत्तरं केवलं संस्कृतभाषायां लिख।\n"
     "६) यदि सन्दर्भे प्रश्नस्य उत्तरं न दृश्यते, तर्हि एवमेव लिख —\n"
     "   \"सन्दर्भे उत्तरं न उपलब्धम्।\"\n"
     "७) सदा उत्तरं शुद्धं संस्कृतं भाषां उपयोजय। कदापि अङ्ग्रेजी वा अन्या भाषा न लिख।"),    
    ("human",
     "सन्दर्भः:\n"
     "{context}\n\n"
     "प्रश्नः:\n"
     "{question}\n\n"
     "उत्तरम् (इतः आरभ्य केवलं उत्तरं लिख):"
    )

    ])

    messages = chat_prompt.format_messages(context=context, question=query)

    prompt_text = "\n".join(m.content for m in messages)

    inputs = tokenizer(
        prompt_text,
        return_tensors="pt",
        truncation=True,
        max_length=512
    )

    input_length = inputs['input_ids'].shape[1]

    with torch.no_grad():
        outputs = model.generate(
            **inputs.to(model.device),
            max_new_tokens=50,
            do_sample=False,
            eos_token_id=tokenizer.eos_token_id,
            pad_token_id=tokenizer.pad_token_id if tokenizer.pad_token_id else tokenizer.eos_token_id
        )

    generated_tokens = outputs[0][input_length:]
    answer_text = tokenizer.decode(generated_tokens, skip_special_tokens=True).strip()

    ai_msg = AIMessage(content=answer_text)

    return ai_msg.content if ai_msg.content else "सन्दर्भे उत्तरं न उपलब्धम्।"


# ============================================================================
# Sidebar Configuration
# ============================================================================

# Hardcoded configuration values
chunk_size = 500
chunk_overlap = 150
num_results = 4

with st.sidebar:
    st.header("📜 Sanskrit RAG System")
    
    st.markdown("""
    ### 🎯 About
    This system uses **Retrieval-Augmented Generation (RAG)** to answer questions 
    based on your Sanskrit documents.
    
    ### ⚙️ Current Settings
    - **Chunk Size:** 500 characters
    - **Chunk Overlap:** 150 characters  
    - **Retrieved Documents:** Top 4 matches
    
    ### 🔧 How It Works
    1. **Upload** your Sanskrit documents (DOCX, PDF, TXT)
    2. **Index** them for semantic search
    3. **Ask** questions in Sanskrit (Devanagari)
    4. **Get** contextual answers from your documents
    
    ### 🚀 Powered By
    - **Embeddings:** Multilingual MiniLM
    - **Search:** FAISS Vector Store
    - **Generation:** Sarvam-1 LLM
    """)
    
    st.markdown("---")
    st.markdown("### 📁 Upload Your Documents")


# ============================================================================
# Initialize Session State
# ============================================================================

if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None
if "retriever" not in st.session_state:
    st.session_state.retriever = None
if "documents" not in st.session_state:
    st.session_state.documents = []
if "tokenizer" not in st.session_state:
    st.session_state.tokenizer = None
if "model" not in st.session_state:
    st.session_state.model = None
if "uploaded_files_info" not in st.session_state:
    st.session_state.uploaded_files_info = []

# ============================================================================
# Load Model (cached)
# ============================================================================

@st.cache_resource
def load_models():
    """Load tokenizer and model (cached)"""
    with st.spinner("Loading Sanskrit generation model..."):
        try:
            tokenizer = AutoTokenizer.from_pretrained("sarvamai/sarvam-1")
            model = AutoModelForCausalLM.from_pretrained(
                "sarvamai/sarvam-1",
                torch_dtype=torch.float32,
                cache_dir=r"models"
            )
            model.to("cpu")
            model.eval()
        except Exception as e:
            st.error(f"Error loading model: {str(e)}")
            raise
    return tokenizer, model


# ============================================================================
# File Upload Section
# ============================================================================

st.header("📂 Upload Documents")
st.markdown("Upload DOCX, PDF, or TXT files to build your knowledge base.")

uploaded_files = st.file_uploader(
    "Choose files",
    type=["docx", "pdf", "txt"],
    accept_multiple_files=True,
    help="Upload multiple documents (DOCX, PDF, TXT)"
)

if uploaded_files:
    if st.button("🔄 Process and Index Documents", key="process_files"):
        with st.spinner("Processing documents..."):
            raw_documents = []
            
            for uploaded_file in uploaded_files:
                # Save uploaded file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                    tmp_file.write(uploaded_file.getvalue())
                    tmp_file_path = tmp_file.name
                
                try:
                    # Extract text
                    file_type = uploaded_file.name.split('.')[-1].lower()
                    text = extract_text_from_file(tmp_file_path, file_type)
                    
                    if text:
                        # Create LangchainDocument objects
                        paragraphs = text.split("\n")
                        for para_num, para in enumerate(paragraphs):
                            if para.strip():
                                langchain_doc = LangchainDocument(
                                    page_content=para,
                                    metadata={
                                        "source": uploaded_file.name,
                                        "file_type": file_type,
                                        "paragraph": para_num
                                    }
                                )
                                raw_documents.append(langchain_doc)
                        
                        st.success(f"✅ Loaded {uploaded_file.name}")
                
                finally:
                    # Clean up temporary file
                    if os.path.exists(tmp_file_path):
                        os.remove(tmp_file_path)
            
            if raw_documents:
                # Process documents
                with st.spinner("Cleaning documents..."):
                    cleaned_docs, num_cleaned = process_documents(raw_documents)
                    st.info(f"📝 Cleaned {num_cleaned} documents")
                
                # Create chunks
                with st.spinner("Creating chunks..."):
                    documents = create_chunks(cleaned_docs)
                    st.info(f"📦 Created {len(documents)} chunks")
                
                # Create vector store
                with st.spinner("Creating embeddings and vector store..."):
                    vectorstore, embeddings = create_vector_store(documents)
                    st.session_state.vectorstore = vectorstore
                    st.session_state.documents = documents
                    st.session_state.retriever = vectorstore.as_retriever(
                        search_type="similarity",
                        search_kwargs={"k": num_results}
                    )
                
                st.session_state.uploaded_files_info = [f.name for f in uploaded_files]
                
                st.success(f"✅ Successfully indexed {len(uploaded_files)} file(s) with {len(documents)} chunks!")
            else:
                st.error("❌ No text could be extracted from the uploaded files.")


# ============================================================================
# Display Indexed Files Info
# ============================================================================

if st.session_state.uploaded_files_info:
    with st.expander("📋 Indexed Files", expanded=False):
        st.markdown("**Files currently in the knowledge base:**")
        for file_name in st.session_state.uploaded_files_info:
            st.markdown(f"- {file_name}")
        st.markdown(f"\n**Total chunks indexed:** {len(st.session_state.documents)}")


# ============================================================================
# Query Section
# ============================================================================

st.header("❓ Ask Questions")

if st.session_state.retriever is None:
    st.warning("⚠️ Please upload and process documents first.")
else:
    # Load models if not already loaded
    if st.session_state.tokenizer is None:
        st.session_state.tokenizer, st.session_state.model = load_models()
    
    # Use form to show button immediately
    with st.form(key="query_form"):
        query = st.text_input(
            "Enter your Sanskrit query (Devanagari script):",
            placeholder="उदाहरणः: देवः कथं साहाय्यम् करोति?",
            help="Query should be in Sanskrit (Devanagari script)"
        )
        submit_button = st.form_submit_button("🔍 Search and Generate Answer")
    
    if submit_button and query:
        # Validate query
        if not is_valid_query(query):
            st.error("❌ Please enter a valid query (minimum 3 characters)")
        else:
                with st.spinner("Retrieving documents and generating answer..."):
                    # Retrieve documents
                    docs = st.session_state.retriever.invoke(query)
                    docs = [d for d in docs if len(d.page_content.strip()) > 10]
                    
                    if not docs:
                        st.warning("⚠️ No relevant documents found in the knowledge base.")
                    else:
                        # Generate answer
                        answer = generate_answer(
                            query,
                            docs,
                            st.session_state.tokenizer,
                            st.session_state.model
                        )
                        
                        # Display results
                        st.markdown("---")
                        col1, col2 = st.columns([1, 1])
                        
                        with col1:
                            st.subheader("❓ Question")
                            st.markdown(f"**{query}**")
                        
                        with col2:
                            st.subheader("✅ Answer")
                            st.markdown(f"**{answer}**")
                        
                        st.markdown("---")
                        
                        # Display retrieved documents
                        with st.expander(f"📚 Retrieved Documents ({len(docs)})", expanded=False):
                            for i, doc in enumerate(docs, 1):
                                st.markdown(f"### Document {i}")
                                st.markdown(f"**Source:** {doc.metadata.get('source', 'Unknown')}")
                                st.markdown(f"**Content:**\n{doc.page_content}")
                                st.markdown("---")


# ============================================================================
# Footer
# ============================================================================

st.markdown("---")
st.markdown("""
### 🔧 Supported File Types
- **DOCX**: Microsoft Word documents
- **PDF**: PDF documents (requires PyPDF2)
- **TXT**: Plain text files

### 📝 Notes
- All queries must be in Sanskrit (Devanagari script)
- The system uses FAISS for efficient similarity search
""")
